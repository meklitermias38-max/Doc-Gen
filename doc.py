from __future__ import annotations

import io
import json
import re
from dataclasses import dataclass
from decimal import Decimal, ROUND_HALF_UP
from typing import Any, Dict, List, Optional, TypedDict, Literal

import streamlit as st
from docx import Document
from docx.shared import Pt

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors
except ImportError:
    SimpleDocTemplate = None

try:
    from google import genai
except ImportError:
    genai = None

try:
    from langgraph.graph import StateGraph, END
except ImportError:
    StateGraph = None
    END = None

try:
    from pydantic import BaseModel, Field, ValidationError
except ImportError:
    BaseModel = object
    Field = None
    ValidationError = Exception


# ============================================================
# CONFIG
# ============================================================

HARDCODED_GEMINI_API_KEY = ""
DEFAULT_MODEL_NAME = "gemini-3.1-pro-preview"

st.set_page_config(page_title="Company Document Generator", layout="wide")
st.title("Company Document Generator")
st.caption("Generate Business Intelligence, Executive Storylines, and ADM using Gemini")


# ============================================================
# DATA MODELS
# ============================================================

@dataclass
class ExecProfile:
    name: str
    title: str
    linkedin: str
    type: str
    business_lines: Optional[List[str]] = None


class AdmFinancialState(TypedDict, total=False):
    company_name: str
    bi_text: str
    consulting_partner: str
    extracted_inputs: Dict[str, Any]
    financial_summary: Dict[str, Any]
    error: str


class CompanyFactsSchema(BaseModel):
    employee_count: int
    annual_revenue_m: float
    sector: str
    legacy_level: str
    scope_preference: str


class BusinessUnitSchema(BaseModel):
    name: str
    estimated_weight_pct: float


class ValueDriverSchema(BaseModel):
    business_unit: str
    driver_name: str
    revenue_or_cost_base_m: float
    improvement_pct: float
    annual_impact_m: float
    source_logic: str


class InvestmentCaseSchema(BaseModel):
    total_investment_m: float
    investment_logic: str
    legacy_pct: float = 59.5
    modernization_pct: float = 15.5
    digital_pods_pct: float = 19.0
    innovation_pct: float = 6.0


class FinancialExtractionSchema(BaseModel):
    company_facts: CompanyFactsSchema
    business_units: List[BusinessUnitSchema]
    value_drivers: List[ValueDriverSchema]
    investment_case: InvestmentCaseSchema
    error: str = ""


class ValidationReportSchema(BaseModel):
    status: Literal["PASSED", "FAILED"]
    financial_errors: List[str] = []
    adm_errors: List[str] = []
    bi_errors: List[str] = []
    warnings: List[str] = []


# ============================================================
# HELPERS
# ============================================================

def sanitize_filename(name: str) -> str:
    return re.sub(r"[^a-zA-Z0-9._ -]+", "", name).strip().replace(" ", "_")


def save_docx_bytes(title: str, body: str) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    doc.add_heading(title, level=0)
    lines = body.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        if line.startswith("|") and "|" in line:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1

            clean_rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.strip("|").split("|")]
                if all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
                    continue
                clean_rows.append(cells)

            if clean_rows:
                table = doc.add_table(rows=1, cols=len(clean_rows[0]))
                table.style = "Table Grid"
                hdr_cells = table.rows[0].cells
                for idx, cell_text in enumerate(clean_rows[0]):
                    hdr_cells[idx].text = cell_text
                    for paragraph in hdr_cells[idx].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                for row in clean_rows[1:]:
                    row_cells = table.add_row().cells
                    for idx, cell_text in enumerate(row[:len(row_cells)]):
                        row_cells[idx].text = cell_text
            continue

        if line.startswith("PART ") or line.startswith("BATCH ") or (line.isupper() and len(line) < 100):
            doc.add_heading(line, level=1)
        elif line.startswith("TABLE ") or re.match(r"^\d+(\.\d+)*\s+", line):
            doc.add_heading(line, level=2)
        elif line.startswith("· ") or line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            p = doc.add_paragraph()
            run = p.add_run(line)
            if line.endswith(":") and len(line) < 100:
                run.bold = True
        i += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def save_pdf_bytes(title: str, body: str) -> bytes:
    if SimpleDocTemplate is None:
        raise ImportError("reportlab is not installed. Run: python -m pip install reportlab")

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    story = [Paragraph(title, styles["Title"]), Spacer(1, 12)]

    lines = body.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            story.append(Spacer(1, 6))
            i += 1
            continue

        if line.startswith("|") and "|" in line:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1

            clean_rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.strip("|").split("|")]
                if all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
                    continue
                clean_rows.append(cells)

            if clean_rows:
                table = Table(clean_rows, repeatRows=1)
                table.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 7),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]))
                story.append(table)
                story.append(Spacer(1, 10))
            continue

        safe_line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if line.startswith("PART ") or line.startswith("BATCH ") or (line.isupper() and len(line) < 100):
            story.append(Paragraph(f"<b>{safe_line}</b>", styles["Heading1"]))
        elif line.startswith("TABLE ") or re.match(r"^\d+(\.\d+)*\s+", line):
            story.append(Paragraph(f"<b>{safe_line}</b>", styles["Heading2"]))
        else:
            story.append(Paragraph(safe_line, styles["BodyText"]))
        i += 1

    doc.build(story)
    buffer.seek(0)
    return buffer.read()



def clean_json_response(text: str) -> str:
    text = text.strip()

    if text.startswith("```"):
        text = text.strip("`")
        if text.lower().startswith("json"):
            text = text[4:].strip()

    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1 and end > start:
        text = text[start:end + 1]

    return text.strip()


def parse_exec_profiles_from_json(raw_json: str) -> List[ExecProfile]:
    cleaned = clean_json_response(raw_json)
    data = json.loads(cleaned)

    items = data.get("executives", [])
    profiles: List[ExecProfile] = []

    for item in items:
        profiles.append(
            ExecProfile(
                name=item.get("name", ""),
                title=item.get("title", ""),
                linkedin=item.get("linkedin", ""),
                type=item.get("type", ""),
                business_lines=item.get("business_lines") or ([item.get("business_line")] if item.get("business_line") else []),
            )
        )

    return profiles


def safe_float(value: Any, default: float = 0.0) -> float:
    try:
        if value is None:
            return default
        if isinstance(value, (int, float)):
            return float(value)
        cleaned = str(value).replace(",", "").replace("$", "").replace("%", "").strip()
        return float(cleaned)
    except Exception:
        return default


def safe_int(value: Any, default: int = 0) -> int:
    try:
        if value is None:
            return default
        if isinstance(value, int):
            return value
        return int(round(safe_float(value, default)))
    except Exception:
        return default


def round1(x: float) -> float:
    return float(Decimal(str(x)).quantize(Decimal("0.1"), rounding=ROUND_HALF_UP))


def mfmt(x: float) -> str:
    return f"${round1(x):,.1f}M"


def pfmt(x: float) -> str:
    return f"{round1(x)}%"


def mfmt_or_na(x: Any) -> str:
    try:
        if x is None:
            return "N/A"
        x = float(x)
        if abs(x) < 0.0001:
            return "N/A"
        return f"${round1(x):,.1f}M"
    except Exception:
        return "N/A"


def pfmt_or_na(x: Any) -> str:
    try:
        if x is None:
            return "N/A"
        x = float(x)
        if abs(x) < 0.0001:
            return "N/A"
        return f"{round1(x)}%"
    except Exception:
        return "N/A"


def contains_bad_zero_values(text: str) -> bool:
    bad_patterns = [
        r"\$0(?:\.0+)?M",
        r"\$0(?:\.0+)?K",
        r"\b0(?:\.0+)?%\b",
        r"\|\s*0(?:\.0+)?\s*\|",
    ]
    return any(re.search(p, text) for p in bad_patterns)


def assert_no_bad_zero_values(text: str, label: str) -> None:
    if contains_bad_zero_values(text):
        raise ValueError(f"{label} contains unsupported zero values. Recalculate and regenerate.")


def allocate_component_total(total: float, weights: List[float]) -> List[float]:
    total = round1(total)
    raw = [total * w for w in weights]
    rounded = [round1(x) for x in raw]
    diff = round1(total - sum(rounded))
    rounded[-1] = round1(rounded[-1] + diff)
    return rounded




def allocate_partner_client_split(total_y: List[float], partner_share: float = 0.42) -> Dict[str, Any]:
    """
    Creates a deterministic partner/client split that reconciles exactly.

    Why this exists:
    - Rounding each year independently can make partner_total fail the 42% test.
    - This function first locks the total split, then allocates the partner share across years.
    - Client yearly values are calculated as the residual, so each year still ties to total_y.
    """
    total_y = [round1(x) for x in total_y]
    total_investment = round1(sum(total_y))

    partner_total, client_total = allocate_component_total(
        total_investment,
        [partner_share, 1.0 - partner_share],
    )

    if total_investment <= 0:
        raise ValueError("Total investment must be positive for partner/client split.")

    year_weights = [y / total_investment for y in total_y]
    partner_y = allocate_component_total(partner_total, year_weights)

    client_y = [round1(total_y[i] - partner_y[i]) for i in range(len(total_y))]
    client_diff = round1(client_total - sum(client_y))
    if client_y:
        client_y[-1] = round1(client_y[-1] + client_diff)

    for i in range(len(total_y)):
        yearly_check = round1(partner_y[i] + client_y[i])
        if yearly_check != total_y[i]:
            client_y[i] = round1(client_y[i] + round1(total_y[i] - yearly_check))

    return {
        "partner_y_m": partner_y,
        "client_y_m": client_y,
        "partner_total_m": round1(sum(partner_y)),
        "client_total_m": round1(sum(client_y)),
    }


def find_payback_years(investment: float, annual_value: float) -> float:
    if investment <= 0 or annual_value <= 0:
        return 0.0

    yearly_values = [
        annual_value * 0.15,
        annual_value * 0.40,
        annual_value * 0.70,
        annual_value * 0.90,
        annual_value * 1.00,
    ]
    cumulative = 0.0
    for idx, yv in enumerate(yearly_values, start=1):
        prev = cumulative
        cumulative += yv
        if cumulative >= investment:
            remaining = investment - prev
            fraction = remaining / yv if yv > 0 else 0
            return round(idx - 1 + fraction, 2)
    return 5.0


def approx_equal(a: float, b: float, tolerance: float = 0.05) -> bool:
    return abs(round1(a) - round1(b)) <= tolerance


def parse_money_values_m(text: str) -> List[float]:
    values = []
    for match in re.findall(r"\$([\d,]+(?:\.\d+)?)\s*M", text):
        values.append(safe_float(match))
    return values


def pydantic_to_dict(model: Any) -> Dict[str, Any]:
    """Works with both Pydantic v1 and v2."""
    if hasattr(model, "model_dump"):
        return model.model_dump()
    if hasattr(model, "dict"):
        return model.dict()
    return dict(model)


def calculate_roi_from_solution(
    five_year_value_m: float,
    investment_m: float,
    annual_maintenance_m: float,
) -> tuple[float, float]:
    """Calculate ROI from actual extracted solution economics. No target ROI band is forced."""
    five_year_value_m = round1(five_year_value_m)
    investment_m = round1(investment_m)
    annual_maintenance_m = round1(annual_maintenance_m)

    if five_year_value_m <= 0:
        raise ValueError("Five-year value must be positive to calculate ROI.")
    if investment_m <= 0:
        raise ValueError("Investment must be positive to calculate ROI.")
    if annual_maintenance_m <= 0:
        raise ValueError("Annual maintenance must be positive to calculate investment multiplier.")

    roi_pct = round1(((five_year_value_m - investment_m) / investment_m) * 100.0)
    multiplier = round1(investment_m / annual_maintenance_m)

    if roi_pct < 100.0:
        raise ValueError(
            f"ROI is below the minimum acceptable threshold. Calculated ROI is {pfmt(roi_pct)}. "
            "Strengthen the value case, reduce the investment scope, or revise phasing based on the actual solution economics."
        )

    return roi_pct, multiplier


def get_proposal_branding(consulting_partner: str) -> Dict[str, str]:
    if consulting_partner == "Direct Client":
        return {
            "proposal_from": "Tholons",
            "prepared_by": "Tholons Inc.",
            "partnership_label": "Tholons-led transformation partnership",
        }

    return {
        "proposal_from": f"{consulting_partner} & Tholons",
        "prepared_by": f"{consulting_partner} & Tholons Inc.",
        "partnership_label": f"{consulting_partner}-Tholons transformation partnership",
    }



# ============================================================
# GEMINI CLIENT
# ============================================================

class GeminiClient:
    def __init__(self, api_key: str, model: str) -> None:
        if genai is None:
            raise ImportError("google-genai is not installed. Run: python -m pip install -U google-genai")
        self.client = genai.Client(api_key=api_key)
        self.model = model

    def generate(self, prompt: str) -> str:
        response = self.client.models.generate_content(
            model=self.model,
            contents=prompt,
        )

        text = getattr(response, "text", None)
        if text and text.strip():
            return text.strip()

        try:
            candidates = getattr(response, "candidates", []) or []
            parts: List[str] = []
            for cand in candidates:
                content = getattr(cand, "content", None)
                if not content:
                    continue
                for part in getattr(content, "parts", []) or []:
                    maybe_text = getattr(part, "text", None)
                    if maybe_text:
                        parts.append(maybe_text)
            return "\n".join(parts).strip()
        except Exception:
            return ""


# ============================================================
# PROMPTS
# ============================================================

BI_PROMPT = """
You are a senior enterprise strategy consultant.

Generate a complete Business Intelligence document for:

Company Name: {company_name}

STRICT RULES
- Identify exactly 3 to 5 top business lines only.
- Business lines must be functional, not geographic.
- If business lines are geography-based, regenerate internally before answering.
- Do not mention or copy any example company.
- Do not use placeholder text.
- Do not include methodology notes.
- Start directly with Section 1.
- Every business line must follow the exact same structure.
- The final summary table must include percentage improvement or ROI-style metric and dollar impact.

REQUIRED OUTPUT STRUCTURE

1. [Business Line Name]

Market Leaders: [List exactly 4 real company names that are actual market leaders or strong benchmark peers for this specific business line and sector. Never write Market Leader A, Market Leader B, leading banks, global peers, regional players, top competitors, or generic descriptions.]

What "Good" Looks Like Today in {company_name}:
· [Bullet 1]
· [Bullet 2]
· [Bullet 3]

What “Good” Looks Like Today Across Market Leaders:
I. [Actual Company Name 1] The "[Benchmark Theme]" Benchmark
· [Specific benchmark explanation]
· [Specific benchmark explanation]

II. [Actual Company Name 2] The "[Benchmark Theme]" Benchmark
· [Specific benchmark explanation]
· [Specific benchmark explanation]

III. [Actual Company Name 3] The "[Benchmark Theme]" Benchmark
· [Specific benchmark explanation]
· [Specific benchmark explanation]

IV. [Actual Company Name 4] The "[Benchmark Theme]" Benchmark
· [Specific benchmark explanation]
· [Specific benchmark explanation]

Challenges faced by {company_name} in [Business Line Name]:
· [Challenge 1]
· [Challenge 2]
· [Challenge 3]

Strategic AI Reinvention and ROI [Business Line Name]: [Theme Name]
Focus: [One focused sentence.]

· Tangible Value/ROI: [Percentage improvement with estimated dollar logic.]
· KYC & Risk Impact or Operational Impact: [Percentage improvement with estimated dollar logic.]
· 5 Daily AI-Driven Nudges:
1. [Nudge 1]
2. [Nudge 2]
3. [Nudge 3]
4. [Nudge 4]
5. [Nudge 5]
· What to do to deliver: [Implementation recommendation.]

Repeat the above structure for each business line.

Summary of Quantified Impact Annual

| Business Unit | Primary Hard ROI Metric | Percentage ROI / Improvement | Estimated Annual Dollar Impact (USD) |
|---|---|---:|---:|
| [Business Unit 1] | [Metric name] | [+X% / -Y%] | $XM - $YM |
| [Business Unit 2] | [Metric name] | [+X% / -Y%] | $XM - $YM |
| [Business Unit 3] | [Metric name] | [+X% / -Y%] | $XM - $YM |

FINAL RULES
- No text after the final summary table.
- The final table must not omit percentage ROI or improvement.
- Do not mention any example company.
- Every Market Leaders line must include exactly 4 actual company names.
- The same 4 named companies must appear again in the "What Good Looks Like Today Across Market Leaders" section.
- Never use placeholders such as Market Leader A, Market Leader B, Competitor 1, Competitor 2, Company A, Company B, Peer A, Peer B, or similar.
- Never use generic descriptions such as leading global banks, regional players, market leaders, global competitors, industry leaders, top players, large incumbents, or benchmark peers instead of company names.
- If a business line is niche, use the closest named public benchmark companies in that sector.
"""

BI_STRUCTURE_FIX_PROMPT = """
You are a strict structure editor.

Rewrite the BI content below so it follows the required structure exactly.
Do not mention or copy any example company.
Do not remove useful quantified details.
Do not add disclaimers.

COMPANY NAME:
{company_name}

REQUIRED STRUCTURE:
1. [Business Line Name]
Market Leaders:
What "Good" Looks Like Today in [Company]:
What “Good” Looks Like Today Across Market Leaders:
Challenges faced by [Company] in [Business Line Name]:
Strategic AI Reinvention and ROI [Business Line Name]:
Summary of Quantified Impact Annual table with columns:
| Business Unit | Primary Hard ROI Metric | Percentage ROI / Improvement | Estimated Annual Dollar Impact (USD) |

BUSINESS INTELLIGENCE DRAFT:
{bi_text}
"""

LEADERSHIP_EXTRACTION_PROMPT = """
You are a strict JSON generator.

Return ONLY valid JSON.
NO explanations. NO markdown. NO text outside JSON.

Schema:
{{
  "executives": [
    {{
      "name": "Full Name",
      "title": "Role Title",
      "linkedin": "",
      "type": "CEO | CFO | CMO | CIO | BUSINESS_LINE_HEAD | BUSINESS_LINE_TECH_HEAD | BOD",
      "business_lines": ["Optional list of business lines this person leads"]
    }}
  ]
}}

Rules:
- Always include "executives"
- If no executives found, return: {{"executives": []}}
- Do not invent names
- Use empty string if linkedin is not provided
- Use BUSINESS_LINE_HEAD for business line leaders
- Use BUSINESS_LINE_TECH_HEAD for technology heads of business lines
- Use BOD for board members
- If one person leads multiple business lines, return all of them in business_lines.

LEADERSHIP MAPPING TEXT:
{leadership_text}
"""

STORYLINE_PROMPT = """
You are a senior strategy advisor creating a high-level executive meeting storyline.

TARGET COMPANY:
{company_name}

BUSINESS INTELLIGENCE:
{business_intelligence}

EXECUTIVE PROFILE:
Name: {name}
Title: {title}
LinkedIn: {linkedin}
Type: {exec_type}
Business Lines Covered: {business_lines}

TASK:
Create a detailed executive storyline customized for this person.

OUTPUT REQUIREMENTS:
- Tailored intro paragraph
- Named meeting storyline theme
- The Hook (Minutes 1-3)
- Proof of Knowledge
- The Pivot (Minutes 4-5)
- Competitive contrast with named peers
- The Close for Action (Minutes 6-7)
- Value proposition with quantified impact
- Detailed meeting structure
- ROI table
"""

FINANCIAL_EXTRACTION_PROMPT = """
You are a strict extraction engine for ADM financial inputs.

The Business Intelligence document below has already been generated.
Do NOT rewrite it. Do NOT summarize it.
Do NOT mention or copy any example company.

Return ONLY valid JSON.
NO markdown.
NO explanations.

Schema:
{{
  "company_facts": {{
    "employee_count": 0,
    "annual_revenue_m": 0,
    "sector": "",
    "legacy_level": "high | moderate | low",
    "scope_preference": "light | medium | heavy"
  }},
  "business_units": [
    {{
      "name": "",
      "estimated_weight_pct": 0
    }}
  ],
  "value_drivers": [
    {{
      "business_unit": "",
      "driver_name": "",
      "revenue_or_cost_base_m": 0,
      "improvement_pct": 0,
      "annual_impact_m": 0,
      "source_logic": ""
    }}
  ],
  "investment_case": {{
    "total_investment_m": 0,
    "investment_logic": "",
    "legacy_pct": 59.5,
    "modernization_pct": 15.5,
    "digital_pods_pct": 19.0,
    "innovation_pct": 6.0
  }},
  "error": ""
}}

Rules:
- annual_revenue_m must be in millions.
- sector must be one of: Financial Services, Semiconductor, Media, Telecom, Manufacturing, Healthcare, Retail.
- Extract all meaningful value drivers from the BI.
- estimated_weight_pct across business_units should sum close to 100.
- total_investment_m must be based on the proposed ADM solution scope, company scale, application portfolio, maintenance baseline, modernization backlog, and delivery ambition.
- Do NOT calculate total_investment_m by forcing ROI to a target number.
- Do NOT target 150%, 180%, 200%, 250%, or 300% ROI.
- ROI must emerge naturally from the value case and investment case.
- The final ROI is allowed to exceed 300% if the underlying value case supports it.
- The only minimum rule is that ROI must not be below 100%.
- Do not invent fake value drivers just to reach a target count.
- Output JSON only.
- Never output 0 for employee_count, annual_revenue_m, revenue_or_cost_base_m, annual_impact_m, improvement_pct, or total_investment_m unless explicitly stated in the BI.
- If a company fact is missing, estimate it from the BI and public scale cues.
- If a value driver is weak or incomplete, derive the annual impact using the best available base and percentage.
- If you cannot support at least 4 value drivers with non-zero annual impact, return an error object instead of zero-filled rows.

COMPANY NAME:
{company_name}

BUSINESS INTELLIGENCE:
{business_intelligence}
"""

NUMERIC_CORRECTION_PROMPT = """
You are a numerical consistency editor.

TASK:
Correct only number discrepancies in the ADM text below using the Financial Summary JSON and the exact tables.
Do NOT rewrite style.
Do NOT change structure.
Do NOT shorten text.
Only replace incorrect numbers, percentages, totals, and repeated values.
Do not add zeros.
Do not add $0.0M or 0.0% placeholders.
Do not mention or copy any example company.

BUSINESS INTELLIGENCE CONTEXT:
{business_intelligence}

FINANCIAL SUMMARY JSON:
{financial_summary_json}

VERBATIM TABLES:
{financial_tables_text}

ADM TEXT TO CORRECT:
{adm_text}

Return only the corrected ADM text.
"""
ADM_BATCH1_PROMPT = """
You are writing BATCH 1 of an ADM proposal.

IMPORTANT SOURCE RULES
- The Business Intelligence document below has already been generated.
- Do not regenerate the BI.
- Do not mention or copy any example company.
- Use BI only as source context.
- The Financial Summary JSON below is the only source of truth for all numbers.
- Do not recalculate, invent, or alter any number.
- Insert the financial tables exactly where relevant.
- If a section does not have support from BI or Financial Summary, omit unsupported zero-value rows.

CLIENT NAME:
{company_name}

CONSULTING PARTNER:
{consulting_partner}

PROPOSAL FROM:
{proposal_from}

PREPARED BY:
{prepared_by}

PARTNERSHIP LABEL:
{partnership_label}

BUSINESS INTELLIGENCE DOCUMENT:
{business_intelligence}

FINANCIAL SUMMARY JSON:
{financial_summary_json}

VERBATIM FINANCIAL TABLES:
{financial_tables_text}

Write BATCH 1 ONLY.

START EXACTLY WITH:
BATCH 1: Writing Executive Summary and Part 1. All numbers from Financial Summary.

THEN WRITE EXACTLY THIS STRUCTURE:

COMPREHENSIVE APPLICATION PORTFOLIO ANALYSIS & 5-YEAR TRANSFORMATION PARTNERSHIP
A Joint Proposal from {proposal_from} to {company_name}

EXECUTIVE SUMMARY: THE STRATEGIC IMPERATIVE

The Executive Summary must be 3 to 5 paragraphs and must include:
- Why the client is at an application modernization inflection point
- Total application count
- Annual maintenance cost
- Technical debt / modernization backlog
- 5-year investment
- ROI
- Cumulative savings
- Annual business value
- Competitive urgency against named market leaders

Then include this exact bullet structure:
{proposal_from} propose a 5-year, [investment] "{partnership_label}" that will:
· Reduce legacy maintenance costs by [percentage/savings] through industrialized offshore delivery and modernization.
· Accelerate digital transformation through API layers, cloud modernization, and AI-enabled delivery.
· Generate [annual value] in annual business value at steady state.
· Transform ADM from a maintenance-heavy cost center into a strategic modernization and innovation engine.

PART 1: DETAILED APPLICATION PORTFOLIO ANALYSIS

1.1 Application Portfolio Composition & Characteristics
Include exactly these three bullets:
· Total Application Count: [number]
· Annual Maintenance Cost: [amount]
· Technical Debt Quantification: [amount] in modernization backlog

Portfolio Distribution by Business Unit:
Use the BUSINESS UNIT ALLOCATION TABLE from the verbatim financial tables.

Technology Stack Distribution:
Use the TECHNOLOGY STACK DISTRIBUTION TABLE from the verbatim financial tables.

For each business unit from the Financial Summary, write one numbered section:
1.2 [Business Unit Name] Deep Dive
1.3 [Business Unit Name] Deep Dive
1.4 [Business Unit Name] Deep Dive, if applicable
1.5 [Business Unit Name] Deep Dive, if applicable
1.6 [Business Unit Name] Deep Dive, if applicable

EACH BUSINESS UNIT DEEP DIVE MUST FOLLOW THIS EXACT STRUCTURE:

[Business Unit Name] Deep Dive

Core Application Systems:

1. [Named System 1]
· Purpose: [one sentence]
· Technology Stack:
· Frontend: [specific technology or plausible stack]
· Middleware: [specific technology or plausible stack]
· Backend: [specific technology or plausible stack]
· Integration: [specific integration pattern]
· Current State Issues:
· [Issue 1 tied to BI]
· [Issue 2 tied to BI]
· [Issue 3 tied to BI]
· Maintenance Cost: [amount from allocation or logical share]
· Market Comparison: [named competitor and relevant benchmark]

2. [Named System 2]
Use the same exact structure.

3. [Named System 3]
Use the same exact structure.

4. [Named System 4]
Use the same exact structure.

If the business unit requires more depth, include:
5. [Named System 5]
Use the same exact structure.

6. [Named System 6]
Use the same exact structure.

Quantifiable Impact — [Business Unit Name]:
Use a table with exactly these columns:
| Business Driver | Per-Unit Metric | Annual Enterprise Impact |
|---|---|---:|

Rules for the Quantifiable Impact table:
- Use value drivers from the Financial Summary.
- Do not invent zero values.
- Do not use $0.0M.
- If only one value driver exists for a business unit, include one row only.
- If multiple value drivers exist, include multiple rows.
- Annual Enterprise Impact must match Financial Summary values.

STRICT FORMAT RULES FOR PART 1:
- Do not use generic headings like "Application Analysis" unless the reference structure requires it.
- Do not summarize business units in paragraphs only.
- Every business unit must have named systems.
- Every named system must include Purpose, Technology Stack, Current State Issues, Maintenance Cost, and Market Comparison.
- Keep the table structures consistent with the reference ADM documents.
- Do not generate Part 2.
- Do not generate Part 3.
- Do not generate Part 4.
- Do not generate Appendices.
- End exactly with: BATCH 1 complete. Say 'continue' for the next batch.
"""
ADM_CONTINUE_PROMPT = """
You are continuing an ADM proposal.

SOURCE RULES
- Do not regenerate the BI.
- Do not mention or copy any example company.
- The Financial Summary JSON is the only source of truth for numbers.
- Do not recalculate, invent, or change numbers.
- Do not add zero-value placeholders.

CLIENT NAME:
{company_name}

CONSULTING PARTNER:
{consulting_partner}

PROPOSAL FROM:
{proposal_from}

PREPARED BY:
{prepared_by}

PARTNERSHIP LABEL:
{partnership_label}

BUSINESS INTELLIGENCE DOCUMENT:
{business_intelligence}

FINANCIAL SUMMARY JSON:
{financial_summary_json}

VERBATIM FINANCIAL TABLES:
{financial_tables_text}

ALREADY GENERATED ADM CONTENT:
{current_adm_text}

Continue with BATCH {next_batch_number} only.

BATCH 2
PART 2: COMPETITIVE BENCHMARKING AGAINST MARKET LEADERS

STRICT STRUCTURE RULES FOR PART 2:
- Each business unit must have exactly one section.
- Each section must contain exactly one benchmarking table.
- Each table must have exactly 4 capability rows.
- Do not use bullet points anywhere in Part 2.
- Do not add paragraphs before the table.
- After each table, add only one sentence beginning with "Quantified Impact:".
- Number the sections as 2.1, 2.2, 2.3, 2.4, and 2.5 if required.

For each business unit, use exactly this format:

Section [number] [Business Unit Name]: [Competitor 1] and [Competitor 2] Benchmark

| Capability | Client Current State | Market Leader Advantage | Technology Gap |
|---|---|---|---|
| [Strategic Capability 1] | [Client current-state issue] | [Named market leader advantage] | [Specific technology gap] |
| [Strategic Capability 2] | [Client current-state issue] | [Named market leader advantage] | [Specific technology gap] |
| [Strategic Capability 3] | [Client current-state issue] | [Named market leader advantage] | [Specific technology gap] |
| [Strategic Capability 4] | [Client current-state issue] | [Named market leader advantage] | [Specific technology gap] |

Quantified Impact: $XXX.XM in annual value at risk if the client fails to close the [specific gap].

BATCH 3
PART 3: 5-YEAR TRANSFORMATION PARTNERSHIP DEAL STRUCTURE
3.1 Partnership Overview & Commercial Terms
3.2 Year-by-Year Investment & Delivery Roadmap
Write Year 1, Year 2, Year 3, Year 4, Year 5 separately.

BATCH 4
3.3 Detailed Financial Model
TABLE A: 5-Year Investment Profile
TABLE B: Business Value Creation
TABLE C: Return on Investment Analysis
3.4 Offshore Delivery Model & Cost Advantage

BATCH 5
3.5 Governance & Operating Model
3.6 Risk Mitigation Framework
3.7 Transition Approach
3.8 Success Metrics & Performance Dashboard

BATCH 6
PART 4: CONCLUSION & STRATEGIC IMPERATIVES
4.1 The Competitive Imperative
4.2 The Partnership Advantage
4.3 Critical Success Factors
4.4 Recommended Next Steps
4.5 Final Investment Summary

APPENDICES

Prepared for:
{company_name} Executive Leadership Team

Prepared by:
{prepared_by}

Date: March 2026

Tholons Contacts:
Abhay Anant Vashistha; abhay@tholons.com
Frank Pendle; frank@tholons.com
Avinash Vashistha; avi@tholons.com

START with:
BATCH {next_batch_number}: Writing requested sections. All numbers from Financial Summary.

END CONDITION:
- If this is Batch 6, end with: BATCH 6 complete. ADM document fully generated.
- If this is not Batch 6, end with: BATCH {next_batch_number} complete. Say 'continue' for the next batch.
"""


# ============================================================
# NUMERICAL AGENT LOGIC
# ============================================================

SECTOR_APP_RATIOS = {
    "Financial Services": 15,
    "Semiconductor": 22,
    "Media": 20,
    "Telecom": 18,
    "Manufacturing": 25,
    "Healthcare": 18,
    "Retail": 22,
}

SECTOR_MAINT_RATIOS = {
    "Financial Services": 0.025,
    "Semiconductor": 0.015,
    "Media": 0.020,
    "Telecom": 0.020,
    "Manufacturing": 0.015,
    "Healthcare": 0.022,
    "Retail": 0.018,
}

LEGACY_MULTIPLIERS = {
    "high": 2.0,
    "moderate": 1.5,
    "low": 1.2,
}


def validate_financial_extraction_with_pydantic(data: Dict[str, Any]) -> FinancialExtractionSchema:
    try:
        parsed = FinancialExtractionSchema(**data)
    except ValidationError as e:
        raise ValueError(f"Pydantic financial extraction validation failed: {e}")

    if parsed.error:
        raise ValueError(parsed.error)

    if parsed.company_facts.employee_count <= 0:
        raise ValueError("Employee count must be non-zero.")
    if parsed.company_facts.annual_revenue_m <= 0:
        raise ValueError("Annual revenue must be non-zero.")
    if parsed.investment_case.total_investment_m <= 0:
        raise ValueError("Total investment must be non-zero and must come from the solution investment case.")
    if len(parsed.value_drivers) < 4:
        raise ValueError("At least 4 non-zero value drivers are required.")

    for driver in parsed.value_drivers:
        if driver.annual_impact_m <= 0:
            raise ValueError(f"Value driver has zero impact: {driver.driver_name}")
        if driver.improvement_pct <= 0:
            raise ValueError(f"Value driver has zero improvement percentage: {driver.driver_name}")

    return parsed


def financial_extract_node(state: AdmFinancialState) -> AdmFinancialState:
    client = st.session_state._gemini_client

    base_prompt = FINANCIAL_EXTRACTION_PROMPT.format(
        company_name=state["company_name"],
        business_intelligence=state["bi_text"],
    )

    raw = client.generate(base_prompt)
    cleaned = clean_json_response(raw)

    try:
        extracted = json.loads(cleaned)
    except Exception as e:
        return {"error": f"JSON parsing failed: {str(e)}\nRaw output:\n{cleaned}"}

    def extracted_has_bad_zeros(data: Dict[str, Any]) -> bool:
        facts = data.get("company_facts", {})
        if safe_float(facts.get("employee_count")) <= 0:
            return True
        if safe_float(facts.get("annual_revenue_m")) <= 0:
            return True
        investment_case = data.get("investment_case", {})
        if safe_float(investment_case.get("total_investment_m")) <= 0:
            return True

        drivers = data.get("value_drivers", [])
        non_zero_drivers = sum(1 for d in drivers if safe_float(d.get("annual_impact_m")) > 0)
        return non_zero_drivers < 4

    if extracted.get("error") or extracted_has_bad_zeros(extracted):
        retry_prompt = base_prompt + """

STRICT RETRY INSTRUCTION:
- Your previous output contained zero, missing, or unusable financial values.
- Recalculate and re-estimate all missing company facts.
- Return at least 4 non-zero value drivers.
- Do not output any driver with annual_impact_m = 0.
- Do not output employee_count = 0, annual_revenue_m = 0, or total_investment_m = 0.
"""
        raw = client.generate(retry_prompt)
        cleaned = clean_json_response(raw)

        try:
            extracted = json.loads(cleaned)
        except Exception as e:
            return {"error": f"Retry JSON parsing failed: {str(e)}\nRaw output:\n{cleaned}"}

    try:
        parsed = validate_financial_extraction_with_pydantic(extracted)
    except Exception as e:
        return {"error": str(e)}

    return {"extracted_inputs": pydantic_to_dict(parsed)}


def build_business_unit_allocations(
    business_units: List[Dict[str, Any]],
    value_drivers: List[Dict[str, Any]],
    app_count: int,
    maintenance_m: float,
    tech_debt_m: float,
) -> List[Dict[str, Any]]:
    impacts_by_unit: Dict[str, float] = {}
    for vd in value_drivers:
        unit = vd.get("business_unit", "Other")
        impacts_by_unit[unit] = impacts_by_unit.get(unit, 0.0) + safe_float(vd.get("annual_impact_m"))

    allocations: List[Dict[str, Any]] = []
    total_weight = 0.0

    for bu in business_units:
        name = bu.get("name", "Business Unit")
        weight = safe_float(bu.get("estimated_weight_pct"))
        if weight <= 0 and name in impacts_by_unit:
            weight = impacts_by_unit[name]
        if weight <= 0:
            weight = 1.0
        allocations.append({"name": name, "weight": weight})
        total_weight += weight

    if total_weight <= 0:
        total_weight = float(len(allocations)) if allocations else 1.0

    for item in allocations:
        pct = item["weight"] / total_weight
        item["portfolio_pct"] = round1(pct * 100)
        item["app_count"] = max(1, int(round(app_count * pct)))
        item["annual_maintenance_m"] = round1(maintenance_m * pct)
        item["modernization_backlog_m"] = round1(tech_debt_m * pct)

    if allocations:
        app_diff = app_count - sum(x["app_count"] for x in allocations)
        allocations[-1]["app_count"] += app_diff

        maint_diff = round1(maintenance_m - sum(x["annual_maintenance_m"] for x in allocations))
        allocations[-1]["annual_maintenance_m"] = round1(allocations[-1]["annual_maintenance_m"] + maint_diff)

        debt_diff = round1(tech_debt_m - sum(x["modernization_backlog_m"] for x in allocations))
        allocations[-1]["modernization_backlog_m"] = round1(allocations[-1]["modernization_backlog_m"] + debt_diff)

        pct_diff = round1(100.0 - sum(x["portfolio_pct"] for x in allocations))
        allocations[-1]["portfolio_pct"] = round1(allocations[-1]["portfolio_pct"] + pct_diff)

    return allocations


def build_blended_rates(sector: str) -> List[Dict[str, Any]]:
    roles = [
        ("Enterprise Architect", 210, 82, "40/60"),
        ("Business Analyst", 155, 62, "35/65"),
        ("Senior Engineer", 175, 68, "25/75"),
        ("Cloud Engineer", 180, 72, "25/75"),
        ("Data Engineer", 160, 60, "20/80"),
        ("Full Stack Developer", 145, 55, "20/80"),
        ("QA Automation Engineer", 120, 45, "10/90"),
        ("Legacy Support Specialist", 110, 40, "10/90"),
    ]

    out: List[Dict[str, Any]] = []
    for role, us, india, mix in roles:
        us_pct, india_pct = mix.split("/")
        us_share = safe_float(us_pct) / 100.0
        india_share = safe_float(india_pct) / 100.0
        blended_exact = (us * us_share) + (india * india_share)
        blended_display = round1(blended_exact)
        savings_pct = round1(((us - blended_exact) / us) * 100)
        out.append(
            {
                "role": role,
                "us_k": us,
                "india_k": india,
                "mix": mix,
                "blended_k": blended_display,
                "savings_pct": savings_pct,
                "formula": f"({us} x {us_share:.2f}) + ({india} x {india_share:.2f}) = {blended_display}",
            }
        )
    return out



def financial_compute_node(state: AdmFinancialState) -> AdmFinancialState:
    if state.get("error"):
        return state

    extracted = state.get("extracted_inputs")
    if not extracted:
        return {"error": "Financial extraction did not produce usable inputs."}

    facts = extracted.get("company_facts", {})
    business_units = extracted.get("business_units", [])
    value_drivers_raw = extracted.get("value_drivers", [])
    investment_case = extracted.get("investment_case", {})

    employee_count = safe_int(facts.get("employee_count"), 0)
    annual_revenue_m = safe_float(facts.get("annual_revenue_m"), 0.0)
    sector = facts.get("sector", "Manufacturing")

    if employee_count <= 0:
        raise ValueError("Employee count was not extracted with a valid non-zero value.")
    if annual_revenue_m <= 0:
        raise ValueError("Annual revenue was not extracted with a valid non-zero value.")
    if sector not in SECTOR_APP_RATIOS:
        sector = "Manufacturing"

    legacy_level = str(facts.get("legacy_level", "moderate")).lower().strip()
    if legacy_level not in LEGACY_MULTIPLIERS:
        legacy_level = "moderate"

    scope_preference = str(facts.get("scope_preference", "medium")).lower().strip()
    if scope_preference not in {"light", "medium", "heavy"}:
        scope_preference = "medium"

    sector_ratio = SECTOR_APP_RATIOS[sector]
    maintenance_ratio = SECTOR_MAINT_RATIOS[sector]
    legacy_multiplier = LEGACY_MULTIPLIERS[legacy_level]

    lower_bound = employee_count / 30
    upper_bound = employee_count / 12
    raw_app_count = employee_count / sector_ratio
    used_ratio = sector_ratio
    if raw_app_count < lower_bound or raw_app_count > upper_bound:
        used_ratio = 20
        raw_app_count = employee_count / used_ratio

    app_count = int(round(raw_app_count))
    annual_maintenance_m = round1(annual_revenue_m * maintenance_ratio)
    tech_debt_m = round1(annual_maintenance_m * legacy_multiplier)

    value_drivers: List[Dict[str, Any]] = []
    for idx, vd in enumerate(value_drivers_raw, start=1):
        bu = vd.get("business_unit", "Business Unit")
        name = vd.get("driver_name", f"Driver {idx}")
        base_m = safe_float(vd.get("revenue_or_cost_base_m"))
        improvement_pct = safe_float(vd.get("improvement_pct"))

        if improvement_pct > 1:
            improvement_decimal = improvement_pct / 100.0
            improvement_pct_display = improvement_pct
        else:
            improvement_decimal = improvement_pct
            improvement_pct_display = improvement_pct * 100.0

        extracted_annual_impact = safe_float(vd.get("annual_impact_m"))
        computed_annual_impact = round1(base_m * improvement_decimal) if base_m > 0 and improvement_decimal > 0 else 0.0
        annual_impact_m = computed_annual_impact if computed_annual_impact > 0 else extracted_annual_impact
        annual_impact_m = round1(annual_impact_m)

        if annual_impact_m <= 0:
            continue

        if base_m > 0 and improvement_decimal > 0:
            formula = f"{round1(base_m)} x {round1(improvement_pct_display)}% = {annual_impact_m}"
        else:
            formula = vd.get("source_logic", "") or f"Annual impact estimated at {annual_impact_m}"

        value_drivers.append(
            {
                "business_unit": bu,
                "driver_name": name,
                "revenue_or_cost_base_m": round1(base_m),
                "improvement_pct": round1(improvement_pct_display),
                "annual_impact_m": annual_impact_m,
                "formula": formula,
            }
        )

    if len(value_drivers) < 4:
        raise ValueError("Not enough valid value drivers were extracted from the BI.")

    total_annual_value_m = round1(sum(v["annual_impact_m"] for v in value_drivers))
    five_year_value_m = round1(total_annual_value_m * 3.15)

    investment_m = round1(safe_float(investment_case.get("total_investment_m")))
    roi_pct, multiplier = calculate_roi_from_solution(
        five_year_value_m=five_year_value_m,
        investment_m=investment_m,
        annual_maintenance_m=annual_maintenance_m,
    )

    cost_savings = {
        "y1_m": round1(annual_maintenance_m * 0.12),
        "y2_m": round1(annual_maintenance_m * 0.22),
        "y3_m": round1(annual_maintenance_m * 0.30),
        "y4_m": round1(annual_maintenance_m * 0.35),
        "y5_m": round1(annual_maintenance_m * 0.38),
    }
    cost_savings["five_year_total_m"] = round1(sum(cost_savings.values()))

    blended_rates = build_blended_rates(sector)

    legacy_pct = safe_float(investment_case.get("legacy_pct"), 59.5) / 100.0
    modernization_pct = safe_float(investment_case.get("modernization_pct"), 15.5) / 100.0
    digital_pct = safe_float(investment_case.get("digital_pods_pct"), 19.0) / 100.0
    innovation_pct = safe_float(investment_case.get("innovation_pct"), 6.0) / 100.0

    component_totals = allocate_component_total(
        investment_m,
        [legacy_pct, modernization_pct, digital_pct, innovation_pct],
    )
    legacy_total = component_totals[0]
    modernization_total = component_totals[1]
    digital_total = component_totals[2]
    innovation_total = component_totals[3]

    legacy_y = allocate_component_total(legacy_total, [0.24, 0.22, 0.20, 0.18, 0.16])
    modernization_y = allocate_component_total(modernization_total, [0.14, 0.22, 0.30, 0.20, 0.14])
    digital_y = allocate_component_total(digital_total, [0.14, 0.22, 0.28, 0.22, 0.14])
    innovation_y = allocate_component_total(innovation_total, [0.10, 0.15, 0.20, 0.25, 0.30])

    total_y = [
        round1(legacy_y[i] + modernization_y[i] + digital_y[i] + innovation_y[i])
        for i in range(5)
    ]
    investment_m = round1(sum(total_y))

    # Recalculate ROI after investment schedule rounding so the final displayed investment and ROI reconcile.
    roi_pct, multiplier = calculate_roi_from_solution(
        five_year_value_m=five_year_value_m,
        investment_m=investment_m,
        annual_maintenance_m=annual_maintenance_m,
    )

    split = allocate_partner_client_split(total_y, partner_share=0.42)
    partner_y = split["partner_y_m"]
    client_y = split["client_y_m"]
    partner_total = split["partner_total_m"]
    client_total = split["client_total_m"]

    partner_margin_low_m = round1(partner_total * 0.18)
    partner_margin_high_m = round1(partner_total * 0.22)

    business_unit_allocations = build_business_unit_allocations(
        business_units=business_units if business_units else [{"name": "Business Unit 1", "estimated_weight_pct": 100}],
        value_drivers=value_drivers,
        app_count=app_count,
        maintenance_m=annual_maintenance_m,
        tech_debt_m=tech_debt_m,
    )

    payback_years = find_payback_years(investment_m, total_annual_value_m)
    annualized_return_pct = round1(roi_pct / 5.0)

    financial_summary = {
        "base_data": {
            "employee_count": employee_count,
            "annual_revenue_m": round1(annual_revenue_m),
            "sector": sector,
            "sector_app_ratio_used": used_ratio,
            "sector_maintenance_ratio_pct": round1(maintenance_ratio * 100),
            "legacy_level": legacy_level,
            "legacy_multiplier": legacy_multiplier,
            "app_count": app_count,
            "annual_maintenance_m": annual_maintenance_m,
            "tech_debt_m": tech_debt_m,
        },
        "business_unit_allocations": business_unit_allocations,
        "value_drivers": value_drivers,
        "total_annual_value_m": total_annual_value_m,
        "five_year_value_m": five_year_value_m,
        "investment_m": investment_m,
        "investment_logic": investment_case.get("investment_logic", "Investment based on proposed ADM solution scope and delivery phasing."),
        "investment_multiplier_used": multiplier,
        "roi_pct": roi_pct,
        "payback_years": payback_years,
        "annualized_return_pct": annualized_return_pct,
        "cost_savings": cost_savings,
        "blended_rates": blended_rates,
        "investment_schedule": {
            "legacy_total_m": legacy_total,
            "modernization_total_m": modernization_total,
            "digital_total_m": digital_total,
            "innovation_total_m": innovation_total,
            "legacy_y_m": legacy_y,
            "modernization_y_m": modernization_y,
            "digital_y_m": digital_y,
            "innovation_y_m": innovation_y,
            "total_y_m": total_y,
        },
        "partner_split": {
            "partner_y_m": partner_y,
            "client_y_m": client_y,
            "partner_total_m": partner_total,
            "client_total_m": client_total,
            "partner_margin_low_m": partner_margin_low_m,
            "partner_margin_high_m": partner_margin_high_m,
        },
    }

    return {"financial_summary": financial_summary}


# ============================================================
# DETERMINISTIC VALIDATORS
# ============================================================

def validate_financial_math(fs: Dict[str, Any]) -> List[str]:
    errors = []

    value_driver_sum = round1(sum(v["annual_impact_m"] for v in fs["value_drivers"]))
    if not approx_equal(value_driver_sum, fs["total_annual_value_m"]):
        errors.append(f"Value drivers sum {mfmt(value_driver_sum)} does not match total annual value {mfmt(fs['total_annual_value_m'])}.")

    expected_5yr = round1(fs["total_annual_value_m"] * 3.15)
    if not approx_equal(expected_5yr, fs["five_year_value_m"]):
        errors.append(f"5-year value should be {mfmt(expected_5yr)}, not {mfmt(fs['five_year_value_m'])}.")

    investment_m = safe_float(fs.get("investment_m"))
    five_year_value_m = safe_float(fs.get("five_year_value_m"))

    if investment_m <= 0:
        errors.append("Investment must be greater than zero.")
    elif five_year_value_m <= 0:
        errors.append("Five-year value must be greater than zero.")
    else:
        expected_roi = round1(((five_year_value_m - investment_m) / investment_m) * 100)
        if not approx_equal(expected_roi, fs["roi_pct"]):
            errors.append(f"ROI should be {pfmt(expected_roi)}, not {pfmt(fs['roi_pct'])}.")
        if fs["roi_pct"] < 100.0:
            errors.append(f"ROI must be at least 100.0%, not {pfmt(fs['roi_pct'])}.")

    multiplier = safe_float(fs.get("investment_multiplier_used"))
    annual_maintenance_m = safe_float(fs["base_data"].get("annual_maintenance_m"))
    expected_multiplier = round1(investment_m / annual_maintenance_m) if investment_m > 0 and annual_maintenance_m > 0 else 0.0
    if multiplier > 0 and not approx_equal(expected_multiplier, multiplier):
        errors.append(f"Investment multiplier should be {expected_multiplier}, not {multiplier}.")
    c = fs["cost_savings"]
    baseline = fs["base_data"]["annual_maintenance_m"]
    fixed = {
        "y1_m": 0.12,
        "y2_m": 0.22,
        "y3_m": 0.30,
        "y4_m": 0.35,
        "y5_m": 0.38,
    }
    for key, pct in fixed.items():
        expected = round1(baseline * pct)
        if not approx_equal(expected, c[key]):
            errors.append(f"Cost savings {key} should be {mfmt(expected)}, not {mfmt(c[key])}.")

    expected_savings_total = round1(c["y1_m"] + c["y2_m"] + c["y3_m"] + c["y4_m"] + c["y5_m"])
    if not approx_equal(expected_savings_total, c["five_year_total_m"]):
        errors.append("5-year cost savings total does not match annual savings sum.")

    s = fs["investment_schedule"]
    component_checks = [
        ("Legacy", s["legacy_y_m"], s["legacy_total_m"]),
        ("Modernization", s["modernization_y_m"], s["modernization_total_m"]),
        ("Digital", s["digital_y_m"], s["digital_total_m"]),
        ("Innovation", s["innovation_y_m"], s["innovation_total_m"]),
    ]
    for label, years, total in component_checks:
        if not approx_equal(sum(years), total):
            errors.append(f"{label} investment row does not sum to its total.")

    yearly_sum = round1(sum(s["total_y_m"]))
    if not approx_equal(yearly_sum, fs["investment_m"]):
        errors.append("Investment schedule yearly total does not equal total investment.")

    for i in range(5):
        col_sum = round1(s["legacy_y_m"][i] + s["modernization_y_m"][i] + s["digital_y_m"][i] + s["innovation_y_m"][i])
        if not approx_equal(col_sum, s["total_y_m"][i]):
            errors.append(f"Investment schedule column Y{i + 1} does not sum correctly.")

    p = fs["partner_split"]
    if not approx_equal(sum(p["partner_y_m"]), p["partner_total_m"]):
        errors.append("Partner yearly split does not sum to partner total.")
    if not approx_equal(sum(p["client_y_m"]), p["client_total_m"]):
        errors.append("Client yearly split does not sum to client total.")

    expected_partner_total = round1(fs["investment_m"] * 0.42)
    if not approx_equal(expected_partner_total, p["partner_total_m"]):
        errors.append("Partner total is not 42% of total investment.")

    expected_client_total = round1(fs["investment_m"] * 0.58)
    if not approx_equal(expected_client_total, p["client_total_m"]):
        errors.append("Client total is not 58% of total investment.")

    for i in range(5):
        split_col = round1(p["partner_y_m"][i] + p["client_y_m"][i])
        expected_col = round1(fs["investment_schedule"]["total_y_m"][i])
        if not approx_equal(split_col, expected_col):
            errors.append(f"Partner/client split column Y{i + 1} does not equal investment schedule total.")

    expected_low = round1(p["partner_total_m"] * 0.18)
    expected_high = round1(p["partner_total_m"] * 0.22)
    if not approx_equal(expected_low, p["partner_margin_low_m"]):
        errors.append("Partner low margin is not 18% of partner revenue.")
    if not approx_equal(expected_high, p["partner_margin_high_m"]):
        errors.append("Partner high margin is not 22% of partner revenue.")

    return errors


def validate_bi_structure(bi_text: str) -> List[str]:
    errors = []
    required = [
        "Market Leaders:",
        'What "Good" Looks Like Today',
        "What “Good” Looks Like Today Across Market Leaders",
        "Challenges faced by",
        "Strategic AI Reinvention and ROI",
        "Summary of Quantified Impact Annual",
        "Percentage ROI / Improvement",
        "Estimated Annual Dollar Impact",
    ]
    for item in required:
        if item not in bi_text:
            errors.append(f"BI missing required structure element: {item}")

    if contains_bad_zero_values(bi_text):
        errors.append("BI contains unsupported zero values.")

    errors.extend(validate_bi_summary_roi_table(bi_text))
    errors.extend(validate_named_market_leaders(bi_text))

    return errors


def validate_bi_summary_roi_table(bi_text: str) -> List[str]:
    errors = []
    if "Summary of Quantified Impact Annual" not in bi_text:
        return ["BI missing Summary of Quantified Impact Annual table."]

    section = bi_text.split("Summary of Quantified Impact Annual", 1)[-1]
    table_lines = [line.strip() for line in section.splitlines() if line.strip().startswith("|")]

    if len(table_lines) < 3:
        return ["BI Summary of Quantified Impact Annual table is missing or malformed."]

    header = table_lines[0]
    if "Percentage ROI / Improvement" not in header:
        errors.append("BI summary table missing Percentage ROI / Improvement column.")

    for line in table_lines[2:]:
        cells = [c.strip() for c in line.strip("|").split("|")]
        if len(cells) < 4:
            errors.append(f"BI summary table row has too few columns: {line}")
            continue
        pct_cell = cells[2]
        if not pct_cell or pct_cell.lower() in {"n/a", "na", "-", "tbd"}:
            errors.append(f"BI summary table has blank ROI/improvement percentage: {line}")
        elif "%" not in pct_cell:
            errors.append(f"BI summary table ROI/improvement cell must include a percentage: {line}")

    return errors


def validate_named_market_leaders(bi_text: str) -> List[str]:
    """
    Validates that each Market Leaders line lists exactly 4 real named companies.

    Important: Do NOT scan the full line for the phrase "market leaders", because
    every valid row starts with "Market Leaders:". Only validate the entries after
    the colon. This avoids false failures for valid rows such as:
    Market Leaders: Novo Nordisk, Sanofi, AstraZeneca, Boehringer Ingelheim
    """
    errors = []

    placeholder_patterns = [
        r"^market leader\s*[a-z0-9]+$",
        r"^competitor\s*[a-z0-9]+$",
        r"^company\s*[a-z0-9]+$",
        r"^peer\s*[a-z0-9]+$",
        r"^player\s*[a-z0-9]+$",
        r"^leader\s*[a-z0-9]+$",
        r"^benchmark peer\s*[a-z0-9]*$",
    ]

    generic_entries = {
        "leading global players",
        "leading global banks",
        "leading banks",
        "regional players",
        "regional banks",
        "top competitors",
        "market leaders",
        "industry leaders",
        "global peers",
        "regional peers",
        "leading companies",
        "large incumbents",
        "benchmark peers",
        "global competitors",
        "regional competitors",
        "top players",
    }

    market_leader_lines = [
        line.strip()
        for line in bi_text.splitlines()
        if line.strip().lower().startswith("market leaders:")
    ]

    if not market_leader_lines:
        return ["BI does not contain any Market Leaders lines."]

    for line in market_leader_lines:
        names_part = line.split(":", 1)[-1].strip()
        names = [x.strip() for x in re.split(r",|;|\|", names_part) if x.strip()]

        if len(names) != 4:
            errors.append(f"Market Leaders must list exactly 4 named companies: {line}")
            continue

        for name in names:
            cleaned = re.sub(r"\s+", " ", name.strip())
            lowered = cleaned.lower()

            if lowered in generic_entries:
                errors.append(f"Market leader entry appears generic instead of named: {name}")
                continue

            for pattern in placeholder_patterns:
                if re.search(pattern, lowered):
                    errors.append(f"Market leader entry appears placeholder-based instead of named: {name}")
                    break

            # A real company name should usually contain letters and not be only a label/number.
            if not re.search(r"[A-Za-z]", cleaned):
                errors.append(f"Invalid market leader name found: {name}")

    return errors

def validate_adm_structure_and_numbers(adm_text: str, fs: Dict[str, Any], adm_batch: int = 0) -> List[str]:
    errors = []

    if contains_bad_zero_values(adm_text):
        errors.append("ADM contains unsupported zero values.")

    required_order = [
        "EXECUTIVE SUMMARY",
        "PART 1: DETAILED APPLICATION PORTFOLIO ANALYSIS",
    ]

    # Validate only the sections that should exist for the current batch.
    # Earlier version expected the full ADM even after Batch 1, so validation failed too early.
    if adm_batch >= 2:
        required_order.append("PART 2: COMPETITIVE BENCHMARKING")

    if adm_batch >= 3:
        required_order.append("PART 3: 5-YEAR TRANSFORMATION PARTNERSHIP DEAL STRUCTURE")

    if adm_batch >= 4:
        required_order.extend([
            "3.3 Detailed Financial Model",
            "3.4 Offshore Delivery Model",
        ])

    if adm_batch >= 5:
        required_order.extend([
            "3.5 Governance",
            "3.6 Risk Mitigation",
            "3.7 Transition",
            "3.8 Success Metrics",
        ])

    if adm_batch >= 6:
        required_order.extend([
            "PART 4: CONCLUSION",
            "APPENDICES",
            "Prepared for:",
        ])

    last_pos = -1
    for section in required_order:
        pos = adm_text.find(section)
        if pos == -1:
            errors.append(f"ADM missing required section: {section}")
        elif pos < last_pos:
            errors.append(f"ADM section appears out of order: {section}")
        else:
            last_pos = pos

    required_numbers = [
        fs["base_data"]["app_count"],
        fs["base_data"]["annual_maintenance_m"],
        fs["base_data"]["tech_debt_m"],
        fs["investment_m"],
        fs["total_annual_value_m"],
        fs["five_year_value_m"],
        fs["roi_pct"],
        fs["cost_savings"]["five_year_total_m"],
    ]

    for num in required_numbers:
        num_str_1 = f"{round1(num):,.1f}"
        num_str_2 = f"{round(num):,}"
        if num_str_1 not in adm_text and num_str_2 not in adm_text:
            errors.append(f"ADM may be missing locked number: {num}")

    expected_total = round1(sum(v["annual_impact_m"] for v in fs["value_drivers"]))
    if not approx_equal(expected_total, fs["total_annual_value_m"]):
        errors.append("Financial summary total annual value does not match driver sum.")

    return errors


def build_validation_report(
    bi_text: str,
    adm_text: str,
    fs: Optional[Dict[str, Any]],
    adm_batch: int = 0,
) -> Dict[str, Any]:
    financial_errors = []
    adm_errors = []
    bi_errors = []
    warnings = []

    bi_errors = validate_bi_structure(bi_text) if bi_text else ["BI has not been generated."]

    if fs:
        financial_errors = validate_financial_math(fs)
    else:
        financial_errors = ["Financial summary has not been generated."]

    if adm_text and fs and adm_batch > 0:
        adm_errors = validate_adm_structure_and_numbers(adm_text, fs, adm_batch=adm_batch)
    elif adm_text and fs and adm_batch <= 0:
        warnings.append("ADM text exists but ADM batch status is 0. Generate ADM Batch 1 again to lock it to the current financial summary.")
    elif adm_text and not fs:
        adm_errors = ["ADM exists but financial summary is missing."]
    else:
        warnings.append("ADM has not been generated yet.")

    status = "PASSED" if not financial_errors and not adm_errors and not bi_errors else "FAILED"

    report = ValidationReportSchema(
        status=status,
        financial_errors=financial_errors,
        adm_errors=adm_errors,
        bi_errors=bi_errors,
        warnings=warnings,
    )
    return pydantic_to_dict(report)


def render_validation_report_text(report: Dict[str, Any]) -> str:
    lines = [f"VALIDATION STATUS: {report['status']}", ""]

    for label, key in [
        ("BI ERRORS", "bi_errors"),
        ("FINANCIAL ERRORS", "financial_errors"),
        ("ADM ERRORS", "adm_errors"),
        ("WARNINGS", "warnings"),
    ]:
        lines.append(label)
        items = report.get(key, [])
        if items:
            for item in items:
                lines.append(f"- {item}")
        else:
            lines.append("- None")
        lines.append("")

    return "\n".join(lines)


def financial_validate_node(state: AdmFinancialState) -> AdmFinancialState:
    if state.get("error"):
        return state

    fs = state.get("financial_summary")
    if not fs:
        return {"error": "Financial summary was not generated."}

    errors = validate_financial_math(fs)

    roi_pct = safe_float(fs.get("roi_pct"))
    if roi_pct < 100.0:
        errors.append(f"ROI must be at least 100.0%. Current ROI is {pfmt(roi_pct)}.")

    return {"error": " | ".join(errors)} if errors else {"error": ""}


def run_financial_graph(client: GeminiClient, company_name: str, bi_text: str) -> Dict[str, Any]:
    if StateGraph is None:
        raise ImportError("langgraph is not installed. Run: python -m pip install -U langgraph langchain-core")

    st.session_state._gemini_client = client

    workflow = StateGraph(AdmFinancialState)
    workflow.add_node("extract_inputs", financial_extract_node)
    workflow.add_node("compute_financials", financial_compute_node)
    workflow.add_node("validate_financials", financial_validate_node)

    workflow.set_entry_point("extract_inputs")
    workflow.add_edge("extract_inputs", "compute_financials")
    workflow.add_edge("compute_financials", "validate_financials")
    workflow.add_edge("validate_financials", END)

    graph = workflow.compile()
    result = graph.invoke({"company_name": company_name, "bi_text": bi_text})

    if result.get("error"):
        raise ValueError(result["error"])

    return result["financial_summary"]


# ============================================================
# TABLE BUILDERS
# ============================================================

def build_table_1_text(fs: Dict[str, Any]) -> str:
    b = fs["base_data"]
    return "\n".join([
        "| Item | Value | Source |",
        "|---|---:|---|",
        f"| Employee Count | {b['employee_count']:,} | Extracted from BI/public basis |",
        f"| Annual Revenue | {mfmt(b['annual_revenue_m'])} | Extracted from BI/public basis |",
        f"| Sector | {b['sector']} | Classification |",
        "",
        "| Metric | Formula | Result |",
        "|---|---|---:|",
        f"| App Count | {b['employee_count']:,} / {b['sector_app_ratio_used']} | {b['app_count']:,} apps |",
        f"| Annual Maintenance | {mfmt(b['annual_revenue_m'])} x {b['sector_maintenance_ratio_pct']}% | {mfmt(b['annual_maintenance_m'])} |",
        f"| Tech Debt | {mfmt(b['annual_maintenance_m'])} x {b['legacy_multiplier']} | {mfmt(b['tech_debt_m'])} |",
    ])


def build_business_unit_allocation_table(fs: Dict[str, Any]) -> str:
    rows = [
        "| Business Unit | App Count | % of Portfolio | Annual Maintenance | Modernization Backlog |",
        "|---|---:|---:|---:|---:|",
    ]
    for bu in fs["business_unit_allocations"]:
        rows.append(
            f"| {bu['name']} | {bu['app_count']} | {bu['portfolio_pct']}% | "
            f"{mfmt(bu['annual_maintenance_m'])} | {mfmt(bu['modernization_backlog_m'])} |"
        )
    rows.append(
        f"| TOTAL | {fs['base_data']['app_count']} | 100% | "
        f"{mfmt(fs['base_data']['annual_maintenance_m'])} | {mfmt(fs['base_data']['tech_debt_m'])} |"
    )
    return "\n".join(rows)


def build_technology_stack_distribution_table(fs: Dict[str, Any]) -> str:
    total_apps = fs["base_data"]["app_count"]
    sector = fs["base_data"]["sector"]

    cats = [
        ("Legacy Core Platforms", 0.28, "15-25", "Critical Shortage", "High"),
        ("Mid-Life Operational Platforms", 0.28, "10-15", "Declining", "Medium-High"),
        ("Modern Cloud / Digital Platforms", 0.18, "3-8", "High Demand", "Low"),
        ("Industry-Specific Systems", 0.16, "8-12", "Available", "Medium"),
        ("SaaS / Enterprise Support Systems", 0.10, "3-5", "Vendor Managed", "Low-Medium"),
    ]

    if sector == "Financial Services":
        cats[0] = ("Legacy Core Banking / Ledger Platforms", 0.30, "15-25", "Critical Shortage", "High")
        cats[3] = ("Risk / Compliance / Treasury Systems", 0.18, "8-12", "Available", "Medium")
    elif sector == "Healthcare":
        cats[0] = ("Legacy Clinical / Commercial Platforms", 0.28, "15-22", "Critical Shortage", "High")
        cats[3] = ("Regulatory / Quality / Manufacturing Systems", 0.18, "8-12", "Available", "Medium")
    elif sector == "Manufacturing":
        cats[0] = ("Legacy Manufacturing / Dealer Platforms", 0.28, "15-25", "Critical Shortage", "High")
        cats[3] = ("Engineering / Product / Quality Systems", 0.18, "8-12", "Available", "Medium")

    counts = [int(round(total_apps * c[1])) for c in cats]
    counts[-1] += total_apps - sum(counts)

    rows = [
        "| Technology Category | App Count | Average Age (Years) | Skills Availability | Risk Level |",
        "|---|---:|---|---|---|",
    ]
    for idx, c in enumerate(cats):
        rows.append(f"| {c[0]} | {counts[idx]} | {c[2]} | {c[3]} | {c[4]} |")
    return "\n".join(rows)


def build_table_2_text(fs: Dict[str, Any]) -> str:
    rows = [
        "| # | Business Unit | Driver Name | Revenue/Cost Base | Improvement % | Annual Impact | Full Formula |",
        "|---|---|---|---:|---:|---:|---|",
    ]
    for i, vd in enumerate(fs["value_drivers"], start=1):
        rows.append(
            f"| {i} | {vd['business_unit']} | {vd['driver_name']} | {mfmt(vd['revenue_or_cost_base_m'])} | "
            f"{vd['improvement_pct']}% | {mfmt(vd['annual_impact_m'])} | {vd['formula']} |"
        )
    rows.append(
        f"| TOTAL ANNUAL VALUE |  |  |  |  | {mfmt(fs['total_annual_value_m'])} | Sum of all drivers |"
    )
    rows.append("")
    rows.append(f"5-Year Value = {mfmt(fs['total_annual_value_m'])} x 3.15 = {mfmt(fs['five_year_value_m'])}")
    rows.append(f"Final Investment = {mfmt(fs['investment_m'])} | ROI = {pfmt(fs['roi_pct'])}")
    return "\n".join(rows)


def build_table_3_text(fs: Dict[str, Any]) -> str:
    c = fs["cost_savings"]
    b = fs["base_data"]["annual_maintenance_m"]
    return "\n".join([
        "| Year | % | Savings | Formula |",
        "|---|---:|---:|---|",
        f"| 1 | 12% | {mfmt(c['y1_m'])} | {mfmt(b)} x 0.12 |",
        f"| 2 | 22% | {mfmt(c['y2_m'])} | {mfmt(b)} x 0.22 |",
        f"| 3 | 30% | {mfmt(c['y3_m'])} | {mfmt(b)} x 0.30 |",
        f"| 4 | 35% | {mfmt(c['y4_m'])} | {mfmt(b)} x 0.35 |",
        f"| 5 | 38% | {mfmt(c['y5_m'])} | {mfmt(b)} x 0.38 |",
        f"| 5-Year Total |  | {mfmt(c['five_year_total_m'])} | Sum |",
    ])


def build_table_4_text(fs: Dict[str, Any]) -> str:
    rows = [
        "| Role | US ($K) | India ($K) | Mix (US/India) | Blended ($K) | Savings% | Formula |",
        "|---|---:|---:|---|---:|---:|---|",
    ]
    for r in fs["blended_rates"]:
        rows.append(
            f"| {r['role']} | {r['us_k']} | {r['india_k']} | {r['mix']} | {r['blended_k']} | {r['savings_pct']}% | {r['formula']} |"
        )
    return "\n".join(rows)


def build_table_5_text(fs: Dict[str, Any]) -> str:
    s = fs["investment_schedule"]
    return "\n".join([
        "| Component | Y1 | Y2 | Y3 | Y4 | Y5 | Total |",
        "|---|---:|---:|---:|---:|---:|---:|",
        f"| Legacy (59.5%) | {mfmt(s['legacy_y_m'][0])} | {mfmt(s['legacy_y_m'][1])} | {mfmt(s['legacy_y_m'][2])} | {mfmt(s['legacy_y_m'][3])} | {mfmt(s['legacy_y_m'][4])} | {mfmt(s['legacy_total_m'])} |",
        f"| Modernization (15.5%) | {mfmt(s['modernization_y_m'][0])} | {mfmt(s['modernization_y_m'][1])} | {mfmt(s['modernization_y_m'][2])} | {mfmt(s['modernization_y_m'][3])} | {mfmt(s['modernization_y_m'][4])} | {mfmt(s['modernization_total_m'])} |",
        f"| Digital Pods (19%) | {mfmt(s['digital_y_m'][0])} | {mfmt(s['digital_y_m'][1])} | {mfmt(s['digital_y_m'][2])} | {mfmt(s['digital_y_m'][3])} | {mfmt(s['digital_y_m'][4])} | {mfmt(s['digital_total_m'])} |",
        f"| Innovation (6%) | {mfmt(s['innovation_y_m'][0])} | {mfmt(s['innovation_y_m'][1])} | {mfmt(s['innovation_y_m'][2])} | {mfmt(s['innovation_y_m'][3])} | {mfmt(s['innovation_y_m'][4])} | {mfmt(s['innovation_total_m'])} |",
        f"| TOTAL | {mfmt(s['total_y_m'][0])} | {mfmt(s['total_y_m'][1])} | {mfmt(s['total_y_m'][2])} | {mfmt(s['total_y_m'][3])} | {mfmt(s['total_y_m'][4])} | {mfmt(fs['investment_m'])} |",
    ])


def build_roi_table_text(fs: Dict[str, Any]) -> str:
    p = fs["partner_split"]
    return "\n".join([
        "| Metric | Value |",
        "|---|---:|",
        f"| Total 5-Year Investment | {mfmt(fs['investment_m'])} |",
        f"| Total 5-Year Value | {mfmt(fs['five_year_value_m'])} |",
        f"| Net Value Created | {mfmt(fs['five_year_value_m'] - fs['investment_m'])} |",
        f"| ROI | {pfmt(fs['roi_pct'])} |",
        f"| Payback Period | {fs['payback_years']} years |",
        f"| Annualized Return | {pfmt(fs['annualized_return_pct'])} |",
        f"| Partner Revenue | {mfmt(p['partner_total_m'])} |",
        f"| Partner Margin Range | {mfmt(p['partner_margin_low_m'])} to {mfmt(p['partner_margin_high_m'])} |",
    ])


def build_business_value_creation_table(fs: Dict[str, Any]) -> str:
    unit_names = [u["name"] for u in fs["business_unit_allocations"]]

    maps = {
        "Revenue Growth": {u: 0.0 for u in unit_names},
        "Cost Reduction": {u: 0.0 for u in unit_names},
        "Risk Mitigation": {u: 0.0 for u in unit_names},
        "Asset Retention": {u: 0.0 for u in unit_names},
    }

    for vd in fs["value_drivers"]:
        unit = vd["business_unit"] if vd["business_unit"] in unit_names else unit_names[0]
        name = vd["driver_name"].lower()
        annual_5yr = round1(vd["annual_impact_m"] * 3.15)

        if any(k in name for k in ["retention", "churn", "renewal", "retained"]):
            maps["Asset Retention"][unit] += annual_5yr
        elif any(k in name for k in ["risk", "delinquency", "compliance", "fraud"]):
            maps["Risk Mitigation"][unit] += annual_5yr
        elif any(k in name for k in ["cost", "efficiency", "savings", "productivity", "cycle", "reduction"]):
            maps["Cost Reduction"][unit] += annual_5yr
        else:
            maps["Revenue Growth"][unit] += annual_5yr

    header = "| Value Driver | " + " | ".join(unit_names) + " | Total |"
    sep = "|---|" + "|".join(["---:" for _ in unit_names]) + "|---:|"
    rows = [header, sep]

    non_zero_labels = []
    for label, mp in maps.items():
        total = round1(sum(mp.values()))
        if total <= 0:
            continue
        non_zero_labels.append(label)
        values = [mfmt_or_na(round1(mp.get(u, 0.0))) for u in unit_names]
        rows.append(f"| {label} | " + " | ".join(values) + f" | {mfmt(total)} |")

    if len(non_zero_labels) > 1:
        total_values = []
        for u in unit_names:
            total_values.append(round1(sum(maps[label].get(u, 0.0) for label in maps)))
        grand_total = round1(sum(total_values))
        rows.append("| TOTAL VALUE | " + " | ".join(mfmt_or_na(v) for v in total_values) + f" | {mfmt(grand_total)} |")

    return "\n".join(rows)



def build_all_financial_tables_text(fs: Dict[str, Any]) -> str:
    chunks = [
        "TABLE 1: BASE DATA\n" + build_table_1_text(fs),
        "BUSINESS UNIT ALLOCATION TABLE\n" + build_business_unit_allocation_table(fs),
        "TECHNOLOGY STACK DISTRIBUTION TABLE\n" + build_technology_stack_distribution_table(fs),
        "TABLE 2: VALUE AND INVESTMENT\n" + build_table_2_text(fs),
        "TABLE 3: COST SAVINGS\n" + build_table_3_text(fs),
        "TABLE 4: BLENDED RATES\n" + build_table_4_text(fs),
        "TABLE 5: INVESTMENT SCHEDULE\n" + build_table_5_text(fs),
        "BUSINESS VALUE CREATION TABLE\n" + build_business_value_creation_table(fs),
        "ROI TABLE\n" + build_roi_table_text(fs),
    ]
    return "\n\n".join(chunks)


def render_financial_summary_text(company_name: str, fs: Dict[str, Any]) -> str:
    text = [
        f"ADM Financial Summary for {company_name}",
        "",
        build_all_financial_tables_text(fs),
        "",
        "Financial summary complete. All deterministic checks passed.",
    ]
    return "\n".join(text)


# ============================================================
# GENERATION FUNCTIONS
# ============================================================

def extract_leadership_json(client: GeminiClient, leadership_text: str) -> str:
    prompt = LEADERSHIP_EXTRACTION_PROMPT.format(leadership_text=leadership_text)
    return client.generate(prompt)


def generate_bi(client: GeminiClient, company_name: str) -> str:
    raw_bi = client.generate(BI_PROMPT.format(company_name=company_name))
    fixed_bi = client.generate(
        BI_STRUCTURE_FIX_PROMPT.format(
            company_name=company_name,
            bi_text=raw_bi,
        )
    )

    errors = validate_bi_structure(fixed_bi)
    if errors:
        retry_prompt = BI_STRUCTURE_FIX_PROMPT.format(
            company_name=company_name,
            bi_text=fixed_bi,
        ) + """

STRICT RETRY:
- Fix every issue found by the BI validator.
- Every Market Leaders section must list exactly 4 actual company names for that specific business line and sector.
- Do not write Market Leader A, Market Leader B, Competitor 1, Company A, Peer A, leading global players, regional competitors, top players, or similar placeholders.
- The same 4 named companies must appear in the benchmark section below.
- The Summary of Quantified Impact Annual table must include a non-blank percentage ROI / improvement value in every row.
- Return the full corrected BI document.
"""
        fixed_bi = client.generate(retry_prompt)

    final_errors = validate_bi_structure(fixed_bi)
    if final_errors:
        raise ValueError("BI validation failed: " + " | ".join(final_errors))

    return fixed_bi



def generate_storylines(
    client: GeminiClient,
    profiles: List[ExecProfile],
    company_name: str,
    business_intelligence: str,
) -> Dict[str, str]:
    results: Dict[str, str] = {}
    progress = st.progress(0)
    total = max(len(profiles), 1)

    for idx, profile in enumerate(profiles, start=1):
        prompt = STORYLINE_PROMPT.format(
            company_name=company_name,
            business_intelligence=business_intelligence,
            name=profile.name,
            title=profile.title,
            linkedin=profile.linkedin,
            exec_type=profile.type,
            business_lines=", ".join(profile.business_lines or []) or "N/A",
        )
        results[f"{profile.type}__{profile.name}"] = client.generate(prompt)
        progress.progress(idx / total)

    return results


def run_numeric_correction(
    client: GeminiClient,
    business_intelligence: str,
    financial_summary: Dict[str, Any],
    financial_tables_text: str,
    adm_text: str,
) -> str:
    prompt = NUMERIC_CORRECTION_PROMPT.format(
        business_intelligence=business_intelligence,
        financial_summary_json=json.dumps(financial_summary, indent=2),
        financial_tables_text=financial_tables_text,
        adm_text=adm_text,
    )
    corrected = client.generate(prompt)
    final_text = corrected if corrected.strip() else adm_text

    if contains_bad_zero_values(final_text):
        retry_prompt = prompt + """

STRICT RETRY:
- The corrected ADM still contains invalid zero values.
- Replace every unsupported zero with the correct number from the Financial Summary JSON.
- If a section is unsupported, remove the zero-filled row rather than leaving $0.0M or 0.0%.
"""
        final_text = client.generate(retry_prompt).strip()

    assert_no_bad_zero_values(final_text, "ADM output")
    return final_text


def generate_adm_batch1(
    client: GeminiClient,
    company_name: str,
    consulting_partner: str,
    business_intelligence: str,
    financial_summary: Dict[str, Any],
    financial_tables_text: str,
) -> str:
    branding = get_proposal_branding(consulting_partner)
    prompt = ADM_BATCH1_PROMPT.format(
        company_name=company_name,
        consulting_partner=consulting_partner,
        proposal_from=branding["proposal_from"],
        prepared_by=branding["prepared_by"],
        partnership_label=branding["partnership_label"],
        business_intelligence=business_intelligence,
        financial_summary_json=json.dumps(financial_summary, indent=2),
        financial_tables_text=financial_tables_text,
    )
    raw = client.generate(prompt)
    corrected = run_numeric_correction(
        client=client,
        business_intelligence=business_intelligence,
        financial_summary=financial_summary,
        financial_tables_text=financial_tables_text,
        adm_text=raw,
    )
    assert_no_bad_zero_values(corrected, "ADM Batch 1")
    return corrected



def generate_adm_next_batch(
    client: GeminiClient,
    company_name: str,
    consulting_partner: str,
    business_intelligence: str,
    financial_summary: Dict[str, Any],
    financial_tables_text: str,
    current_adm_text: str,
    next_batch_number: int,
) -> str:
    branding = get_proposal_branding(consulting_partner)
    prompt = ADM_CONTINUE_PROMPT.format(
        company_name=company_name,
        consulting_partner=consulting_partner,
        proposal_from=branding["proposal_from"],
        prepared_by=branding["prepared_by"],
        partnership_label=branding["partnership_label"],
        business_intelligence=business_intelligence,
        financial_summary_json=json.dumps(financial_summary, indent=2),
        financial_tables_text=financial_tables_text,
        current_adm_text=current_adm_text,
        next_batch_number=next_batch_number,
    )
    raw = client.generate(prompt)
    corrected = run_numeric_correction(
        client=client,
        business_intelligence=business_intelligence,
        financial_summary=financial_summary,
        financial_tables_text=financial_tables_text,
        adm_text=raw,
    )
    assert_no_bad_zero_values(corrected, f"ADM Batch {next_batch_number}")
    return corrected



# ============================================================
# SESSION STATE
# ============================================================

defaults = {
    "leadership_json": "",
    "bi_text": "",
    "storylines": {},
    "financial_summary": None,
    "financial_summary_text": "",
    "financial_tables_text": "",
    "adm_text": "",
    "adm_batch": 0,
    "validation_report": {},
    "validation_report_text": "",
    "consulting_partner": "Deloitte",
}

for key, value in defaults.items():
    if key not in st.session_state:
        st.session_state[key] = value


# ============================================================
# SIDEBAR
# ============================================================

st.sidebar.header("Configuration")

api_key = HARDCODED_GEMINI_API_KEY.strip()
if not api_key:
    api_key = st.secrets.get("GEMINI_API_KEY", "")
if not api_key:
    api_key = st.sidebar.text_input("Gemini API Key", type="password")

model_name = DEFAULT_MODEL_NAME
st.sidebar.text_input("Model", value=model_name, disabled=True)

company_name = st.sidebar.text_input("Target Company Name")
consulting_partner = st.sidebar.selectbox(
    "Consulting Partner",
    ["Deloitte", "APEX", "UST", "Direct Client"],
    index=["Deloitte", "APEX", "UST", "Direct Client"].index(st.session_state.get("consulting_partner", "Deloitte")),
)
st.session_state.consulting_partner = consulting_partner

st.sidebar.markdown("### Leadership Mapping Text")
leadership_text = st.sidebar.text_area(
    "Paste Leadership Mapping Text",
    height=350,
    placeholder="Paste CEO, CFO, CMO, CIO, business line heads, business line tech heads, and board members here..."
)


# ============================================================
# VALIDATION
# ============================================================

def validate_base() -> bool:
    if not api_key:
        st.error("Enter your Gemini API key.")
        return False
    if not company_name:
        st.error("Enter a target company name.")
        return False
    return True


def validate_leadership() -> bool:
    if not leadership_text.strip():
        st.error("Paste the leadership mapping text first.")
        return False
    return True


def validate_bi() -> bool:
    if not st.session_state.bi_text:
        st.error("Generate Business Intelligence first.")
        return False
    return True


def validate_financial_summary() -> bool:
    if not st.session_state.financial_summary:
        st.error("Generate the ADM Financial Summary first.")
        return False
    return True


def refresh_validation_report() -> None:
    report = build_validation_report(
        bi_text=st.session_state.bi_text,
        adm_text=st.session_state.adm_text,
        fs=st.session_state.financial_summary,
        adm_batch=st.session_state.get("adm_batch", 0),
    )
    st.session_state.validation_report = report
    st.session_state.validation_report_text = render_validation_report_text(report)


def clear_downstream_outputs(clear_bi: bool = False) -> None:
    """
    Prevent stale outputs from being validated against newly generated upstream data.
    Example: if a new financial summary is generated, the old ADM must be cleared
    because it was created from old locked numbers.
    """
    if clear_bi:
        st.session_state.bi_text = ""

    st.session_state.storylines = {}
    st.session_state.financial_summary = None
    st.session_state.financial_summary_text = ""
    st.session_state.financial_tables_text = ""
    st.session_state.adm_text = ""
    st.session_state.adm_batch = 0
    st.session_state.validation_report = {}
    st.session_state.validation_report_text = ""


def clear_adm_outputs_only() -> None:
    """Clear ADM after financial numbers change, while preserving BI and financial summary."""
    st.session_state.adm_text = ""
    st.session_state.adm_batch = 0
    st.session_state.validation_report = {}
    st.session_state.validation_report_text = ""


# ============================================================
# UI
# ============================================================

st.subheader("Inputs")
col1, col2, col3 = st.columns(3)

with col1:
    st.markdown("**Company**")
    st.write(company_name or "Not set")

with col2:
    st.markdown("**Consulting Partner**")
    st.write(consulting_partner)

with col3:
    st.markdown("**Model**")
    st.write(model_name)

st.markdown("**Leadership Mapping Status**")
if leadership_text.strip():
    st.success("Leadership mapping text loaded.")
else:
    st.info("No leadership mapping text pasted yet.")

st.markdown("**ADM Numerical Agent Status**")
if StateGraph is None:
    st.warning("LangGraph not installed yet. Install with: python -m pip install -U langgraph langchain-core")
else:
    st.success("LangGraph available. ADM financial agent ready.")

col_a, col_b, col_c, col_d, col_e, col_f = st.columns(6)

with col_a:
    bi_btn = st.button("Generate BI", use_container_width=True)

with col_b:
    storyline_btn = st.button("Generate Storylines", use_container_width=True)

with col_c:
    financial_btn = st.button("Generate Financial Summary", use_container_width=True)

with col_d:
    adm_batch1_btn = st.button("Generate ADM Batch 1", use_container_width=True)

with col_e:
    adm_continue_btn = st.button("Continue ADM", use_container_width=True)

with col_f:
    validate_btn = st.button("Validate", use_container_width=True)


# ============================================================
# ACTIONS
# ============================================================

if bi_btn:
    if validate_base():
        try:
            client = GeminiClient(api_key=api_key, model=model_name)
            with st.spinner("Generating Business Intelligence..."):
                new_bi_text = generate_bi(client, company_name)
                clear_downstream_outputs(clear_bi=False)
                st.session_state.bi_text = new_bi_text
            refresh_validation_report()
            st.success("Business Intelligence generated. Downstream ADM and financial outputs were reset to prevent stale-number validation.")
        except Exception as e:
            st.error(f"BI generation failed: {e}")

if storyline_btn:
    if validate_base() and validate_leadership() and validate_bi():
        try:
            client = GeminiClient(api_key=api_key, model=model_name)

            with st.spinner("Extracting leadership structure from text..."):
                leadership_json = extract_leadership_json(client, leadership_text)
                st.session_state.leadership_json = leadership_json

            st.subheader("Leadership Extraction Debug")
            st.code(leadership_json)

            profiles = parse_exec_profiles_from_json(leadership_json)

            with st.spinner("Generating executive storylines..."):
                st.session_state.storylines = generate_storylines(
                    client=client,
                    profiles=profiles,
                    company_name=company_name,
                    business_intelligence=st.session_state.bi_text,
                )

            st.success(f"Executive storylines generated for {len(profiles)} profiles.")
        except Exception as e:
            st.error(f"Storyline generation failed: {e}")
            if st.session_state.leadership_json:
                st.subheader("Leadership Extraction Debug")
                st.code(st.session_state.leadership_json)

if financial_btn:
    if validate_base() and validate_bi():
        try:
            client = GeminiClient(api_key=api_key, model=model_name)
            with st.spinner("Running LangGraph ADM financial agent and deterministic validator..."):
                financial_summary = run_financial_graph(
                    client=client,
                    company_name=company_name,
                    bi_text=st.session_state.bi_text,
                )
                st.session_state.financial_summary = financial_summary
                st.session_state.financial_tables_text = build_all_financial_tables_text(financial_summary)
                st.session_state.financial_summary_text = render_financial_summary_text(company_name, financial_summary)
                clear_adm_outputs_only()
            refresh_validation_report()
            st.success("ADM Financial Summary generated and validated. Existing ADM was reset because locked numbers changed.")
        except Exception as e:
            st.error(f"Financial summary generation failed: {e}")

if adm_batch1_btn:
    if validate_base() and validate_bi():
        try:
            if not st.session_state.financial_summary:
                client = GeminiClient(api_key=api_key, model=model_name)
                with st.spinner("Financial summary not found. Running LangGraph ADM financial agent first..."):
                    financial_summary = run_financial_graph(
                        client=client,
                        company_name=company_name,
                        bi_text=st.session_state.bi_text,
                    )
                    st.session_state.financial_summary = financial_summary
                    st.session_state.financial_tables_text = build_all_financial_tables_text(financial_summary)
                    st.session_state.financial_summary_text = render_financial_summary_text(company_name, financial_summary)

            client = GeminiClient(api_key=api_key, model=model_name)
            with st.spinner("Generating ADM Batch 1 using locked numbers..."):
                batch1_text = generate_adm_batch1(
                    client=client,
                    company_name=company_name,
                    consulting_partner=consulting_partner,
                    business_intelligence=st.session_state.bi_text,
                    financial_summary=st.session_state.financial_summary,
                    financial_tables_text=st.session_state.financial_tables_text,
                )
                st.session_state.adm_text = st.session_state.financial_summary_text + "\n\n" + batch1_text
                st.session_state.adm_batch = 1
            refresh_validation_report()
            st.success("ADM Financial Summary + Batch 1 generated.")
        except Exception as e:
            st.error(f"ADM batch 1 generation failed: {e}")

if adm_continue_btn:
    if validate_base() and validate_bi() and validate_financial_summary():
        if st.session_state.adm_batch <= 0:
            st.error("Generate ADM Batch 1 first.")
        elif st.session_state.adm_batch >= 6:
            st.info("ADM is already complete.")
        else:
            try:
                client = GeminiClient(api_key=api_key, model=model_name)
                next_batch = st.session_state.adm_batch + 1
                with st.spinner(f"Generating ADM Batch {next_batch} using locked numbers..."):
                    new_batch_text = generate_adm_next_batch(
                        client=client,
                        company_name=company_name,
                        consulting_partner=consulting_partner,
                        business_intelligence=st.session_state.bi_text,
                        financial_summary=st.session_state.financial_summary,
                        financial_tables_text=st.session_state.financial_tables_text,
                        current_adm_text=st.session_state.adm_text,
                        next_batch_number=next_batch,
                    )
                    st.session_state.adm_text += "\n\n" + new_batch_text
                    st.session_state.adm_batch = next_batch
                refresh_validation_report()
                st.success(f"ADM Batch {next_batch} generated.")
            except Exception as e:
                st.error(f"ADM continuation failed: {e}")

if validate_btn:
    refresh_validation_report()
    if st.session_state.validation_report.get("status") == "PASSED":
        st.success("Validation passed.")
    else:
        st.error("Validation failed. Check the Validation Report tab.")


# ============================================================
# OUTPUTS
# ============================================================

st.divider()
st.header("Outputs")

tab1, tab2, tab3, tab4, tab5 = st.tabs(
    ["Business Intelligence", "Executive Storylines", "ADM Financial Summary", "ADM", "Validation Report"]
)

with tab1:
    st.text_area("Business Intelligence Output", value=st.session_state.bi_text, height=500)

    if st.session_state.bi_text:
        st.download_button(
            "Download BI TXT",
            data=st.session_state.bi_text,
            file_name=f"{sanitize_filename(company_name)}_Business_Intelligence.txt",
            mime="text/plain",
            use_container_width=True,
        )
        st.download_button(
            "Download BI DOCX",
            data=save_docx_bytes(f"{company_name} Business Intelligence", st.session_state.bi_text),
            file_name=f"{sanitize_filename(company_name)}_Business_Intelligence.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
        st.download_button(
            "Download BI PDF",
            data=save_pdf_bytes(f"{company_name} Business Intelligence", st.session_state.bi_text),
            file_name=f"{sanitize_filename(company_name)}_Business_Intelligence.pdf",
            mime="application/pdf",
            use_container_width=True,
        )

with tab2:
    if st.session_state.storylines:
        selected = st.selectbox("Select Storyline", list(st.session_state.storylines.keys()))
        selected_text = st.session_state.storylines[selected]

        st.text_area("Storyline Output", value=selected_text, height=500)

        st.download_button(
            "Download Selected Storyline TXT",
            data=selected_text,
            file_name=f"{sanitize_filename(selected)}.txt",
            mime="text/plain",
            use_container_width=True,
        )
        st.download_button(
            "Download Selected Storyline DOCX",
            data=save_docx_bytes(selected, selected_text),
            file_name=f"{sanitize_filename(selected)}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

        combined = []
        for k, v in st.session_state.storylines.items():
            combined.append(f"\n{'=' * 80}\n{k}\n{'=' * 80}\n{v}\n")
        combined_text = "\n".join(combined)

        st.download_button(
            "Download All Storylines TXT",
            data=combined_text,
            file_name=f"{sanitize_filename(company_name)}_Executive_Storylines.txt",
            mime="text/plain",
            use_container_width=True,
        )
        st.download_button(
            "Download All Storylines DOCX",
            data=save_docx_bytes(f"{company_name} Executive Storylines", combined_text),
            file_name=f"{sanitize_filename(company_name)}_Executive_Storylines.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    else:
        st.info("No storylines generated yet.")

with tab3:
    st.text_area("ADM Financial Summary Output", value=st.session_state.financial_summary_text, height=650)

    if st.session_state.financial_summary_text:
        st.download_button(
            "Download Financial Summary TXT",
            data=st.session_state.financial_summary_text,
            file_name=f"{sanitize_filename(company_name)}_ADM_Financial_Summary.txt",
            mime="text/plain",
            use_container_width=True,
        )
        st.download_button(
            "Download Financial Summary DOCX",
            data=save_docx_bytes(f"{company_name} ADM Financial Summary", st.session_state.financial_summary_text),
            file_name=f"{sanitize_filename(company_name)}_ADM_Financial_Summary.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
        st.download_button(
            "Download Financial Summary PDF",
            data=save_pdf_bytes(f"{company_name} ADM Financial Summary", st.session_state.financial_summary_text),
            file_name=f"{sanitize_filename(company_name)}_ADM_Financial_Summary.pdf",
            mime="application/pdf",
            use_container_width=True,
        )

with tab4:
    st.markdown(f"**ADM Progress:** Batch {st.session_state.adm_batch} / 6")
    st.text_area("ADM Output", value=st.session_state.adm_text, height=700)

    if st.session_state.adm_text:
        st.download_button(
            "Download ADM TXT",
            data=st.session_state.adm_text,
            file_name=f"{sanitize_filename(company_name)}_ADM.txt",
            mime="text/plain",
            use_container_width=True,
        )
        st.download_button(
            "Download ADM DOCX",
            data=save_docx_bytes(f"{company_name} ADM", st.session_state.adm_text),
            file_name=f"{sanitize_filename(company_name)}_ADM.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
        st.download_button(
            "Download ADM PDF",
            data=save_pdf_bytes(f"{company_name} ADM", st.session_state.adm_text),
            file_name=f"{sanitize_filename(company_name)}_ADM.pdf",
            mime="application/pdf",
            use_container_width=True,
        )

with tab5:
    if not st.session_state.validation_report_text:
        refresh_validation_report()

    status = st.session_state.validation_report.get("status", "FAILED")
    if status == "PASSED":
        st.success("Validation Status: PASSED")
    else:
        st.error("Validation Status: FAILED")

    st.text_area("Validation Report", value=st.session_state.validation_report_text, height=500)

    if st.session_state.validation_report_text:
        st.download_button(
            "Download Validation Report TXT",
            data=st.session_state.validation_report_text,
            file_name=f"{sanitize_filename(company_name)}_Validation_Report.txt",
            mime="text/plain",
            use_container_width=True,
        )
